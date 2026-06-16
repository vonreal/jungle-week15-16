from __future__ import annotations

from io import BytesIO
import json
import re
from pathlib import Path

from app.core.config import settings

TECH_KEYWORDS = {
    "언어": {
        "Java": ["java"],
        "Python": ["python", "파이썬"],
        "JavaScript": ["javascript", "java script", "자바스크립트"],
        "TypeScript": ["typescript", "type script", "타입스크립트"],
        "C": ["c언어", " c ", " c,", " c.", " c/"],
        "C++": ["c++", "cpp"],
        "Swift": ["swift"],
        "Kotlin": ["kotlin", "코틀린"],
        "Dart": ["dart"],
        "SQL": ["sql"],
    },
    "백엔드": {
        "Spring": ["spring", "스프링"],
        "Spring Boot": ["spring boot", "스프링 부트"],
        "FastAPI": ["fastapi", "fast api"],
        "Node.js": ["node.js", "node js", "nodejs"],
        "Django": ["django"],
        "Flask": ["flask"],
        "JPA": ["jpa"],
        "Redis": ["redis", "레디스"],
        "Kafka": ["kafka", "카프카"],
        "PostgreSQL": ["postgresql", "postgres"],
        "MySQL": ["mysql"],
        "SQLite": ["sqlite"],
        "MongoDB": ["mongodb", "mongo db"],
    },
    "프론트엔드": {
        "React": ["react", "리액트"],
        "Vue": ["vue"],
        "HTML": ["html"],
        "CSS": ["css"],
        "Tailwind CSS": ["tailwind"],
    },
    "모바일": {
        "Flutter": ["flutter", "플러터"],
        "iOS": ["ios"],
        "Android": ["android", "안드로이드"],
        "UIKit": ["uikit"],
        "Storyboard": ["storyboard", "스토리보드"],
        "Realm": ["realm"],
        "WebSocket": ["websocket", "web socket", "웹소켓"],
        "WebRTC": ["webrtc", "web rtc"],
    },
    "AI": {
        "OpenAI API": ["openai api", "openai"],
        "LangChain": ["langchain", "랭체인"],
        "LangGraph": ["langgraph", "랭그래프"],
        "RAG": ["rag", "검색 증강", "검색증강"],
    },
    "운영/협업": {
        "Git": ["git"],
        "GitHub": ["github", "깃허브"],
        "GitLab": ["gitlab", "깃랩"],
        "Docker": ["docker", "도커"],
        "Kubernetes": ["kubernetes", "쿠버네티스", "k8s"],
        "AWS": ["aws"],
        "Nginx": ["nginx"],
        "Linux": ["linux"],
        "Unix": ["unix"],
        "Cron": ["cron", "크론"],
        "CI/CD": ["ci/cd", "cicd", "ci cd"],
    },
}


class DocumentTextExtractorService:
    def extract_text(self, filename: str | None, body: bytes) -> str:
        suffix = Path(filename or "").suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf(body)
        if suffix == ".docx":
            return self._extract_docx(body)
        return body.decode("utf-8", errors="ignore")

    def _extract_pdf(self, body: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(body))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(page.strip() for page in pages if page.strip())

    def _extract_docx(self, body: bytes) -> str:
        from docx import Document

        document = Document(BytesIO(body))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        table_cells = [
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        return "\n".join(text for text in [*paragraphs, *table_cells] if text)


class DocumentParserService:
    async def extract_experiences(self, raw_text: str | None) -> list[str]:
        if not raw_text:
            return []

        cleaned_text = self._remove_reference_sections(raw_text)
        llm_experiences = await self._extract_with_llm(cleaned_text)
        if llm_experiences:
            return llm_experiences[:20]

        return self._extract_with_rules(cleaned_text)[:20]

    async def extract_skill_mentions(self, raw_text: str | None) -> list[dict[str, str]]:
        if not raw_text:
            return []

        cleaned_text = self._remove_reference_sections(raw_text)
        keyword_mentions = self._extract_skill_mentions_with_keywords(cleaned_text)
        llm_mentions = await self._extract_skill_mentions_with_llm(cleaned_text)
        return self._dedupe_skill_mentions([*keyword_mentions, *llm_mentions])

    def _extract_skill_mentions_with_keywords(self, cleaned_text: str) -> list[dict[str, str]]:
        compact_text = re.sub(r"\s+", " ", cleaned_text).lower()
        normalized = f" {compact_text} "
        mentions: list[dict[str, str]] = []
        seen: set[str] = set()
        for category, skills in TECH_KEYWORDS.items():
            for name, aliases in skills.items():
                if name.lower() in seen:
                    continue
                if any(self._contains_skill_alias(normalized, alias) for alias in aliases):
                    mentions.append(
                        {
                            "category": category,
                            "name": name,
                            "description": "이력서/포트폴리오에서 자동 감지된 기술입니다. 숙련도를 직접 확인해 조정하세요.",
                        }
                    )
                    seen.add(name.lower())
        return mentions

    async def _extract_skill_mentions_with_llm(self, raw_text: str) -> list[dict[str, str]]:
        if not settings.openai_api_key:
            return []

        prompt = (
            "아래 이력서/포트폴리오 텍스트에서 실제 업무나 프로젝트에 사용된 기술 스택만 추출하세요.\n"
            "숙련도 기준표, 단순 기술 목록, 연락처, 학력, 회사 소개는 제외하세요.\n"
            "category는 언어, 백엔드, 프론트엔드, 모바일, AI, 운영/협업 중 하나로 분류하세요.\n"
            "너무 일반적인 단어(개발, 프로젝트, API, 서버, 데이터)는 제외하세요.\n"
            "최대 12개만 반환하고 반드시 JSON 배열만 반환하세요.\n"
            "예: [{\"category\":\"모바일\",\"name\":\"Riverpod\",\"description\":\"문서에서 자동 감지된 기술입니다.\"}]\n\n"
            f"텍스트:\n{raw_text[:12000]}"
        )

        try:
            from langchain_openai import ChatOpenAI

            llm = ChatOpenAI(model=settings.llm_model, api_key=settings.openai_api_key, temperature=0)
            response = await llm.ainvoke(prompt)
        except Exception:
            return []

        return self._parse_skill_mentions_json(str(response.content))

    async def _extract_with_llm(self, raw_text: str) -> list[str]:
        if not settings.openai_api_key:
            return []

        prompt = (
            "아래 이력서/포트폴리오 텍스트에서 실제 프로젝트, 업무, 개발 활동 경험만 추출하세요.\n"
            "숙련도 기준표, 기술 스택 표, 레벨 설명, 단순 기술 목록, 링크 제목, 연락처, 학력은 제외하세요.\n"
            "줄바꿈으로 갈라진 한 경험은 하나의 자연스러운 한국어 문장으로 병합하세요.\n"
            "각 항목은 25~180자 사이로 작성하고, 최대 15개만 반환하세요.\n"
            "반드시 JSON 문자열 배열만 반환하세요. 예: [\"경험 1\", \"경험 2\"]\n\n"
            f"텍스트:\n{raw_text[:12000]}"
        )

        try:
            from langchain_openai import ChatOpenAI

            llm = ChatOpenAI(model=settings.llm_model, api_key=settings.openai_api_key, temperature=0)
            response = await llm.ainvoke(prompt)
        except Exception:
            return []

        return self._parse_experience_json(str(response.content))

    def _parse_experience_json(self, content: str) -> list[str]:
        match = re.search(r"\[[\s\S]*\]", content)
        if not match:
            return []
        try:
            parsed = json.loads(match.group(0))
        except json.JSONDecodeError:
            return []
        if not isinstance(parsed, list):
            return []
        return self._dedupe_experiences(str(item) for item in parsed)

    def _parse_skill_mentions_json(self, content: str) -> list[dict[str, str]]:
        match = re.search(r"\[[\s\S]*\]", content)
        if not match:
            return []
        try:
            parsed = json.loads(match.group(0))
        except json.JSONDecodeError:
            return []
        if not isinstance(parsed, list):
            return []

        mentions: list[dict[str, str]] = []
        for item in parsed:
            if not isinstance(item, dict):
                continue
            name = re.sub(r"\s+", " ", str(item.get("name", ""))).strip()
            category = re.sub(r"\s+", " ", str(item.get("category", ""))).strip()
            if not self._is_skill_name_candidate(name):
                continue
            if category not in {"언어", "백엔드", "프론트엔드", "모바일", "AI", "운영/협업"}:
                category = "운영/협업"
            mentions.append(
                {
                    "category": category,
                    "name": name[:120],
                    "description": "이력서/포트폴리오에서 AI가 자동 감지한 기술입니다. 숙련도를 직접 확인해 조정하세요.",
                }
            )
        return mentions

    def _extract_with_rules(self, raw_text: str) -> list[str]:
        lines = [re.sub(r"\s+", " ", line).strip() for line in raw_text.splitlines()]
        lines = [line for line in lines if line]
        candidates: list[str] = []
        buffer = ""
        for line in lines:
            if self._is_noise_line(line):
                continue
            has_bullet = bool(re.match(r"^[•●*-]\s*", line))
            starts_new = has_bullet or (not buffer and self._looks_like_experience_start(line))
            cleaned = re.sub(r"^[•●*-]\s*", "", line).strip()
            if starts_new and buffer:
                candidates.append(buffer)
                buffer = cleaned
            elif buffer:
                buffer = f"{buffer} {cleaned}".strip()
            else:
                buffer = cleaned
        if buffer:
            candidates.append(buffer)

        experiences = []
        for candidate in candidates:
            cleaned = re.sub(r"\s+", " ", candidate).strip()
            if self._is_experience_candidate(cleaned):
                experiences.append(cleaned)

        return self._dedupe_experiences(experiences)

    def _remove_reference_sections(self, raw_text: str) -> str:
        lines = raw_text.splitlines()
        filtered: list[str] = []
        skipping = False
        for line in lines:
            normalized = re.sub(r"\s+", " ", line).strip()
            if re.search(r"(기본기|언어 및 도구|도구) 숙련도|숙련도 기준|기술역량|주요 분야", normalized):
                skipping = True
                continue
            if skipping and re.match(r"^\s*\d{1,2}\.\s+", normalized):
                skipping = False
            if skipping:
                continue
            filtered.append(line)
        return "\n".join(filtered)

    def _looks_like_experience_start(self, line: str) -> bool:
        return bool(
            re.search(
                r"(프로젝트|개발|구현|설계|운영|배포|개선|분석|마이그레이션|자동화|연동|구축)",
                line,
            )
        )

    def _is_noise_line(self, line: str) -> bool:
        lower = line.lower()
        if len(line) < 8:
            return True
        if re.search(r"^\d+\s*[-–]\s*(연습|초보|기본|실무)\s*숙련도", line):
            return True
        if re.search(r"^(주요 분야|기술역량|수준|업무역량|프로그래밍 언어|개발 프레임워크|운영환경구축)", line):
            return True
        if re.search(r"\b(slack|git|java|javascript|python|swift|spring|node js|nginx|sqlite|unix|linux)\s*\(\d\)", lower):
            return True
        if re.search(r"숙련도|수업 수준|기술구현|사용되는지는 이해하지 못함|장애\(에러\)", line):
            return True
        if re.search(r"^\S+\s+\S+\s+\d$", line):
            return True
        return False

    def _is_experience_candidate(self, text: str) -> bool:
        lower = text.lower()
        if len(text) < 18 or len(text) > 700:
            return False
        if self._is_noise_line(text):
            return False
        keywords = ["개발", "구현", "설계", "운영", "배포", "개선", "분석", "api", "자동화", "연동", "구축"]
        return any(key in lower for key in keywords)

    def _contains_skill_alias(self, normalized_text: str, alias: str) -> bool:
        normalized_alias = alias.lower().strip()
        if re.fullmatch(r"[a-z0-9.+#/-]+", normalized_alias):
            return bool(
                re.search(
                    rf"(?<![a-z0-9.+#/-]){re.escape(normalized_alias)}(?![a-z0-9.+#/-])",
                    normalized_text,
                )
            )
        return normalized_alias in normalized_text

    def _is_skill_name_candidate(self, name: str) -> bool:
        if len(name) < 2 or len(name) > 60:
            return False
        lowered = name.lower()
        blocked = {
            "api",
            "server",
            "frontend",
            "backend",
            "project",
            "개발",
            "서버",
            "프로젝트",
            "데이터",
            "기술",
        }
        if lowered in blocked:
            return False
        return bool(re.search(r"[A-Za-z가-힣0-9+#./-]", name))

    def _dedupe_skill_mentions(self, mentions: list[dict[str, str]]) -> list[dict[str, str]]:
        results: list[dict[str, str]] = []
        seen: set[str] = set()
        for mention in mentions:
            name = re.sub(r"\s+", " ", mention["name"]).strip()
            key = name.lower()
            if key in seen or not self._is_skill_name_candidate(name):
                continue
            seen.add(key)
            results.append({**mention, "name": name})
        return results[:20]

    def _dedupe_experiences(self, experiences) -> list[str]:
        results: list[str] = []
        seen: set[str] = set()
        for experience in experiences:
            cleaned = re.sub(r"\s+", " ", str(experience)).strip(" -•●\n\t")
            if not cleaned or self._is_noise_line(cleaned):
                continue
            key = cleaned.lower()
            if key in seen:
                continue
            seen.add(key)
            results.append(cleaned)
        return results
